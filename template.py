import xlrd
import re
from docx import Document
import locale
from itertools import takewhile
import time

locale.setlocale(locale.LC_TIME, '')  # Heure en Français

dateActuelle = time.strftime('%d %B %Y')

numberLines = 0;
document = Document(r"C:\Users\ppintus\Documents\csd\courrier.docx")


wb = xlrd.open_workbook(r"C:\Users\ppintus\Documents\csd\\testexcel.xls");

sh = wb.sheet_by_name('Feuil1')

nombreLigne = (len(sh.col_values(0)))
nombreColonne = sh.row_len(0);



def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)



regex = re.compile(r"champ1");
replace = str(int(round((sh.row_values(1)[0]))))
docx_replace_regex(document, regex, replace);


regex = re.compile(r"champ3");
replace = str(sh.row_values(1)[3]);
docx_replace_regex(document, regex, replace);







cell_type = sh.cell_type(5, 3)
if (cell_type == xlrd.XL_CELL_EMPTY):
    print("Cette cellule est blanche")
else:
    print(cell_type)


for i in range(1, nombreLigne):
    if (type(sh.row_values(i)[1]) is str):
        replace = sh.row_values(i)[1]
    else:
        replace = int(sh.row_values(i)[1])
        replace = str(replace)

    docx_replace_regex(document, regex, replace)
    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest" + str(i) + "." + "docx")

    regex = re.compile(r"champ2")
    replace = dateActuelle;
    docx_replace_regex(document, regex, replace)
    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest"+str(i)+"."+"docx")

    regex = re.compile(r"champ3");
    replace = str(sh.row_values(i)[3]);
    docx_replace_regex(document, regex, replace);

    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest"+str(i)+"."+"docx")

    regex = re.compile(r"champ4")
    if (type(sh.row_values(i)[4]) is str):
        replace = sh.row_values(i)[4]
    else:
        replace=int(sh.row_values(i)[4])
        replace=str(replace)

    docx_replace_regex(document, regex, replace)
    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest"+str(i)+"."+"docx")


    regex =re.compile(r"champ5")
    replace = sh.row_values(i)[5];
    replace = str(replace)
    docx_replace_regex(document,regex,replace)
    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest"+str(i)+"."+"docx")



    regex = re.compile(r"champ6")
    if (type(sh.row_values(i)[6]) is str):
        replace = sh.row_values(i)[6]
    else:
        replace=int(sh.row_values(i)[6])
        replace=str(replace)
    docx_replace_regex(document,regex,replace)
    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest"+str(i)+"."+"docx")



    docx_replace_regex(document, regex, replace)
    document.save(r"C:\Users\ppintus\Documents\csd\courrierTest"+str(i)+"."+"docx")



    document =Document(r"C:\Users\ppintus\Documents\csd\courrier.docx");





# document.save(r"C:\Users\ppintus\Documents\csd\courrier.docx")


# regex1 = re.compile(r"Caverne")
# replace1 = r"Monsieur le président"
# docx_replace_regex(document, regex1 , replace1)


# docx_replace_regex((document,regex,replace1));
# document.save(r"C:\Users\ppintus\Documents\csd\courrier1.docx")


