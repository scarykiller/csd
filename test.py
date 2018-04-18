import xlrd
import re
from docx import Document

numberLines=0;

document = Document(r"C:\Users\ppintus\Documents\csd\courrier.docx")

wb = xlrd.open_workbook(r"C:\Users\ppintus\Documents\csd\\testexcel.xls");
print(wb.sheet_names());
sh= wb.sheet_by_name('Feuil1')
for rownum in range(1,sh.nrows):
    a=(sh.row_values(rownum));
    print(a[0])
    print(a[1])
    print(a[2])

replace1=re.compile((sh.row_values(1)[0]).toString());
#CHANGER EN STRING

regex=r"champ1";
document.save(r"C:\Users\ppintus\Documents\csd\courrier.docx")




def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

regex1 = re.compile(r"Caverne")
replace1 = r"Monsieur le président"
docx_replace_regex(document, regex1 , replace1)
